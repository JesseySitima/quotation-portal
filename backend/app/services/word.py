from io import BytesIO
from docx import Document


def create_quotation_word(quotation, items):
    document = Document()

    document.add_heading("Quotation Request", level=1)

    document.add_paragraph(
        f"Request Number: {quotation['request_number']}"
    )

    document.add_paragraph(
        f"Facility: {quotation['facility_name']}"
    )

    document.add_paragraph(
        f"Contact Person: {quotation['contact_person']}"
    )

    document.add_paragraph(
        f"Email: {quotation['email']}"
    )

    document.add_paragraph(
        f"Phone: {quotation['phone']}"
    )

    document.add_heading("Requested Items", level=2)

    table = document.add_table(
        rows=1,
        cols=3,
    )

    table.style = "Table Grid"

    header = table.rows[0].cells

    header[0].text = "Product"
    header[1].text = "Quantity"
    header[2].text = "Unit"

    for item in items:
        row = table.add_row().cells

        row[0].text = str(item["product_name"])
        row[1].text = str(item["quantity"])
        row[2].text = str(item["unit"])

    output = BytesIO()

    document.save(output)

    output.seek(0)

    return output.getvalue()